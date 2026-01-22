import os
import PyPDF2
import pdfplumber
from pdf2image import convert_from_path
import base64
from io import BytesIO
from PIL import Image, ImageDraw, ImageFont
import fitz  # PyMuPDF
import img2pdf
from docx import Document
import pandas as pd

class PDFProcessor:
    def __init__(self):
        self.temp_dir = "temp"
        os.makedirs(self.temp_dir, exist_ok=True)
    
    def get_page_count(self, pdf_path):
        """Get total number of pages in PDF"""
        with open(pdf_path, 'rb') as file:
            pdf = PyPDF2.PdfReader(file)
            return len(pdf.pages)
    
    def get_page_as_image(self, pdf_path, page_number, dpi=150):
        """Convert PDF page to base64 image for display"""
        try:
            # Convert PDF page to image
            images = convert_from_path(pdf_path, first_page=page_number, last_page=page_number, dpi=dpi)
            
            if images:
                # Convert to base64
                buffered = BytesIO()
                images[0].save(buffered, format="PNG", optimize=True)
                img_str = base64.b64encode(buffered.getvalue()).decode()
                
                return f"data:image/png;base64,{img_str}"
        except Exception as e:
            print(f"Error converting page to image: {e}")
            return None
    
    def add_text_to_pdf(self, input_path, output_path, text, page_num=1, x=50, y=50):
        """Add text to PDF at specified position"""
        try:
            # Open PDF
            doc = fitz.open(input_path)
            page = doc[page_num - 1]
            
            # Add text
            page.insert_text((x, y), text, fontsize=11, color=(0, 0, 0))
            
            # Save
            doc.save(output_path)
            doc.close()
            return True
        except Exception as e:
            print(f"Error adding text: {e}")
            return False
    
    def highlight_text(self, input_path, output_path, page_num, text, color="#ffff00"):
        """Highlight text in PDF"""
        try:
            doc = fitz.open(input_path)
            page = doc[page_num - 1]
            
            # Search for text
            text_instances = page.search_for(text)
            
            # Highlight each occurrence
            for inst in text_instances:
                highlight = page.add_highlight_annot(inst)
                highlight.set_colors(stroke=fitz.utils.getColor(color))
                highlight.update()
            
            doc.save(output_path)
            doc.close()
            return True
        except Exception as e:
            print(f"Error highlighting text: {e}")
            return False
    
    def add_image_to_pdf(self, input_path, output_path, image_data, page_num, x=50, y=50):
        """Add image to PDF"""
        try:
            doc = fitz.open(input_path)
            page = doc[page_num - 1]
            
            # Decode base64 image
            if image_data.startswith('data:image'):
                image_data = image_data.split(',')[1]
            
            img_bytes = base64.b64decode(image_data)
            
            # Add image
            rect = fitz.Rect(x, y, x + 100, y + 100)  # Adjust size as needed
            page.insert_image(rect, stream=img_bytes)
            
            doc.save(output_path)
            doc.close()
            return True
        except Exception as e:
            print(f"Error adding image: {e}")
            return False
    
    def merge_pdfs(self, pdf_paths, output_path):
        """Merge multiple PDFs"""
        try:
            merger = PyPDF2.PdfMerger()
            
            for pdf_path in pdf_paths:
                merger.append(pdf_path)
            
            merger.write(output_path)
            merger.close()
            return True
        except Exception as e:
            print(f"Error merging PDFs: {e}")
            return False
    
    def split_pdf(self, input_path, pages):
        """Split PDF into multiple files"""
        try:
            output_files = []
            with open(input_path, 'rb') as file:
                pdf = PyPDF2.PdfReader(file)
                
                for i, page_range in enumerate(pages):
                    writer = PyPDF2.PdfWriter()
                    
                    # Handle page ranges like "1-3" or single pages "2"
                    if '-' in str(page_range):
                        start, end = map(int, page_range.split('-'))
                        for page_num in range(start-1, end):
                            writer.add_page(pdf.pages[page_num])
                    else:
                        writer.add_page(pdf.pages[int(page_range)-1])
                    
                    output_filename = f"{os.path.splitext(input_path)[0]}_part_{i+1}.pdf"
                    with open(output_filename, 'wb') as output_file:
                        writer.write(output_file)
                    
                    output_files.append(output_filename)
            
            return output_files
        except Exception as e:
            print(f"Error splitting PDF: {e}")
            return []
    
    def compress_pdf(self, input_path, quality='medium'):
        """Compress PDF file size"""
        try:
            # Simple compression by optimizing images
            doc = fitz.open(input_path)
            
            output_path = input_path.replace('.pdf', '_compressed.pdf')
            
            # Set compression parameters
            compress_level = {
                'low': 1,
                'medium': 2,
                'high': 3
            }.get(quality, 2)
            
            doc.save(output_path, 
                    garbage=4,  # Remove unused objects
                    deflate=True,  # Compress streams
                    deflate_images=True,
                    deflate_fonts=True,
                    compression_level=compress_level)
            
            doc.close()
            return output_path
        except Exception as e:
            print(f"Error compressing PDF: {e}")
            return input_path
    
    def extract_text(self, pdf_path):
        """Extract all text from PDF"""
        try:
            text = ""
            with pdfplumber.open(pdf_path) as pdf:
                for page in pdf.pages:
                    text += page.extract_text() + "\n"
            return text
        except Exception as e:
            print(f"Error extracting text: {e}")
            return ""
    
    def convert_to_docx(self, pdf_path):
        """Convert PDF to Word document"""
        try:
            text = self.extract_text(pdf_path)
            
            doc = Document()
            doc.add_heading('Converted PDF Document', 0)
            
            for paragraph in text.split('\n'):
                if paragraph.strip():
                    doc.add_paragraph(paragraph)
            
            output_path = pdf_path.replace('.pdf', '.docx')
            doc.save(output_path)
            return output_path
        except Exception as e:
            print(f"Error converting to DOCX: {e}")
            return None
    
    def convert_to_excel(self, pdf_path):
        """Convert PDF tables to Excel"""
        try:
            tables = []
            with pdfplumber.open(pdf_path) as pdf:
                for page in pdf.pages:
                    page_tables = page.extract_tables()
                    for table in page_tables:
                        if table:
                            tables.append(table)
            
            if tables:
                # Combine all tables or save separately
                df = pd.DataFrame(tables[0])
                output_path = pdf_path.replace('.pdf', '.xlsx')
                df.to_excel(output_path, index=False)
                return output_path
            else:
                return None
        except Exception as e:
            print(f"Error converting to Excel: {e}")
            return None
    
    def convert_to_images(self, pdf_path, dpi=150):
        """Convert PDF pages to images"""
        try:
            images = convert_from_path(pdf_path, dpi=dpi)
            output_paths = []
            
            for i, image in enumerate(images):
                output_path = f"{pdf_path.replace('.pdf', '')}_page_{i+1}.png"
                image.save(output_path, 'PNG')
                output_paths.append(output_path)
            
            return output_paths
        except Exception as e:
            print(f"Error converting to images: {e}")
            return []

# Singleton instance
pdf_processor = PDFProcessor()

# Convenience functions
get_page_count = pdf_processor.get_page_count
get_page_as_image = pdf_processor.get_page_as_image
add_text_to_pdf = pdf_processor.add_text_to_pdf
highlight_text = pdf_processor.highlight_text
add_image_to_pdf = pdf_processor.add_image_to_pdf
merge_pdfs = pdf_processor.merge_pdfs
split_pdf = pdf_processor.split_pdf
compress_pdf = pdf_processor.compress_pdf
extract_text = pdf_processor.extract_text
convert_to_docx = pdf_processor.convert_to_docx
convert_to_excel = pdf_processor.convert_to_excel
convert_to_images = pdf_processor.convert_to_images